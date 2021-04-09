from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException

options = webdriver.ChromeOptions()
options.add_argument("--start-maximized")
prefs = {"profile.default_content_settings.popups": 0,
             "download.default_directory":
                        r"C:\Users\thegi\PycharmProjects\ViasatAutomation\\",
             "directory_upgrade": True}
options.add_experimental_option("prefs", prefs)

driver = webdriver.Chrome(executable_path=r'C:\Users\thegi\PycharmProjects\ViasatAutomation\chromedriver.exe', options = options)


# Tech notes
#Tim most recent April 9th
#Kendall most recent April 7th




invoice_name = "Jim"

last_check_date = "04/08/2021"

todays_date = "04/08/2021"

invoice_name_date = "JamesCrabtree4921"


from selenium.webdriver.common.keys import Keys
import time

driver.maximize_window()
driver.get("https://fulfillment.wildblue.net")
time.sleep(5)

#x = input("user name")
#y = input("password")

elem = driver.find_element_by_name('j_username')
elem.send_keys("michael.gibson.pv")
elem = driver.find_element_by_name('j_password')
elem.send_keys("TeeHiggins85!")
elem = driver.find_element_by_name('submit')
elem.click()

driver.get("https://fulfillment.wildblue.net/fsm-fe/fsm/browseOrder/browseOrder.page?execution=e3s1")

try:
    elem = driver.find_element_by_id("browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt196") or driver.find_element_by_id("browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt196")
    elem.click()
except NoSuchElementException:
    elem = driver.find_element_by_id("browseOrder:orderGrid:filterBoard:createdDateFilter:j_idt189")
    elem.click()




elem = driver.find_element_by_id("browseOrder:orderGrid:filterBoard:scheduledDateFilter:dateFromCal_input")
elem.click()
elem.send_keys(last_check_date)
elem = driver.find_element_by_id("browseOrder:orderGrid:filterBoard:scheduledDateFilter:dateToCal_input")
elem.click()
elem.send_keys(todays_date)
elem = driver.find_element_by_id('browseOrder:orderGrid:search')
elem.click()
time.sleep(17)
elem = driver.find_element_by_name('browseOrder:xlsExport')
elem.click()
driver.minimize_window()
time.sleep(17)

import os

path = r'C:\Users\thegi\PycharmProjects\ViasatAutomation\\'

files = []
# r=root, d=directories, f = files
for r, d, f in os.walk(path):
    for file in f:
        if '.XLS' in file:
            files.append(os.path.join(r, file))

for f in files:
    print(f)

import pandas as pd
from pandas import ExcelWriter
from pandas import ExcelFile
import fpdf
pdf = fpdf.FPDF(format='letter')
pdf.add_page()
pdf.set_font("Arial", size=12)

df = pd.read_excel(f)



pdf.write(5, invoice_name_date)
pdf.ln()

c = (len(df))
x = 0

complete_date = []
job_names = []
order_type = []
order_region = []
job_pay = []

def Make_Invoice(tech):
    x = 0
    for num in range(c):
        a = (df.iloc[x, 27])
        if a == tech:
            complete_date.append(df.iloc[x, 6])
            job_names.append(df.iloc[x, 14])
            order_type.append(df.iloc[x, 8])
            order_region.append(df.iloc[x, 29])
            if (df.iloc[x, 8]) == "Install":
                if (df.iloc[x, 29]) == "NC09":
                    job_pay.append(150)
                elif (df.iloc[x, 29]) == "NC08":
                    job_pay.append(150)
                elif (df.iloc[x, 29]) == "NC36":
                    job_pay.append(150)
                elif (df.iloc[x, 29]) == "NC03":
                    job_pay.append(150)
                else:
                    job_pay.append(125)
            elif (df.iloc[x, 8]) == "Service Call":
                job_pay.append(65)
            elif (df.iloc[x, 8]) == "Upgrade":
                job_pay.append(85)
            elif (df.iloc[x, 8]) == "Performance Follow Up":
                job_pay.append(65)
            elif (df.iloc[x, 8]) == "Commercial Install":
                job_pay.append(200)
            elif (df.iloc[x, 8]) == "Commercial Service Call":
                job_pay.append(100)
            elif (df.iloc[x, 8]) == "Equipment Change":
                job_pay.append(65)
            elif (df.iloc[x, 8]) == "Commercial Performance Follow Up":
                job_pay.append(100)
        x = x +1



Make_Invoice(invoice_name)



import docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture('viasat.jpg', width=Inches(1.25))

document.add_heading(invoice_name_date, 0)

z = len(job_names) + 1
q = z -1
y = 0
x = 1

table = document.add_table(rows=z, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Complete Date'
hdr_cells[1].text = 'Customer Name'
hdr_cells[2].text = 'Job Type'
hdr_cells[3].text = 'Job Pay'
for row in range(q):
    info_cells = table.rows[x].cells
    info_cells[0].text = str(complete_date[y])
    info_cells[1].text = job_names[y]
    info_cells[2].text = order_type[y]
    info_cells[3].text = str(job_pay[y])
    y = y + 1
    x = x + 1

total = sum(job_pay)

document.add_paragraph(f'Invoice Total                  {total}')


document.add_paragraph('DEBITS:')

document.save(f'{invoice_name_date}.docx')

os.system(f'start {invoice_name_date}.docx')

os.remove(f)



